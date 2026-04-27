"""
Instruction Parser — Parse .docx, .rtfd, .pdf, .txt, .zip instruction files.

Extracts raw text from instruction documents for task extraction.
Supports structured SAM instruction parsing with step splitting,
context propagation, and cross-reference resolution.
"""

from __future__ import annotations

import logging
import re
import subprocess
import zipfile
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class InstructionStep:
    """A single parsed instruction step."""
    step_number: int
    text: str
    sheet_context: Optional[str] = None
    parent_step: Optional[int] = None


class InstructionParser:
    """Parse instruction files of various formats into plain text."""

    SUPPORTED_EXTENSIONS = {".docx", ".rtfd", ".pdf", ".txt", ".rtf", ".zip"}

    def parse(self, path: Path) -> str:
        """
        Parse an instruction file and return its text content.
        Dispatches to the appropriate parser based on extension.
        """
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(f"Instruction file not found: {path}")

        ext = path.suffix.lower()
        if path.is_dir() and str(path).endswith(".rtfd"):
            ext = ".rtfd"

        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported format '{ext}'. Supported: {self.SUPPORTED_EXTENSIONS}"
            )

        parser_map = {
            ".docx": self._parse_docx,
            ".rtfd": self._parse_rtfd,
            ".rtf": self._parse_rtf,
            ".pdf": self._parse_pdf,
            ".txt": self._parse_txt,
            ".zip": self._parse_zip,
        }

        text = parser_map[ext](path)
        text = self._clean_text(text)
        logger.info("Parsed %s: %d characters", path.name, len(text))
        return text

    def parse_multiple(self, paths: list[Path]) -> str:
        """Parse multiple instruction files and concatenate their text."""
        texts = []
        for p in paths:
            try:
                texts.append(self.parse(p))
            except Exception as e:
                logger.warning("Failed to parse %s: %s", p, e)
        return "\n\n".join(texts)

    # ── Format-specific parsers ──

    def _parse_docx(self, path: Path) -> str:
        """Parse a .docx file using python-docx or textutil fallback."""
        try:
            import docx
            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            return "\n".join(paragraphs)
        except ImportError:
            logger.info("python-docx not available, falling back to textutil")
            return self._textutil_convert(path)

    def _parse_rtfd(self, path: Path) -> str:
        """
        Parse an .rtfd bundle (macOS rich text with attachments).
        Uses textutil to convert to plain text.
        """
        return self._textutil_convert(path)

    def _parse_rtf(self, path: Path) -> str:
        """Parse an .rtf file using textutil."""
        return self._textutil_convert(path)

    def _parse_pdf(self, path: Path) -> str:
        """Parse a PDF file using pdfplumber, PyPDF2, or macOS textutil."""
        # Try pdfplumber first (best extraction)
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        texts.append(text)
            if texts:
                return "\n\n".join(texts)
        except ImportError:
            pass

        # Try PyPDF2
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            texts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)
            if texts:
                return "\n\n".join(texts)
        except ImportError:
            pass

        # Fallback: macOS mdimport / textutil
        try:
            result = subprocess.run(
                ["mdimport", "-d2", str(path)],
                capture_output=True, text=True, timeout=30,
            )
            if result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass

        raise RuntimeError(
            f"Cannot parse PDF '{path.name}': install pdfplumber or PyPDF2"
        )

    def _parse_txt(self, path: Path) -> str:
        """Parse a plain text file."""
        return path.read_text(encoding="utf-8", errors="replace")

    def _parse_zip(self, path: Path) -> str:
        """Extract and parse instruction files from a ZIP archive.

        Looks for supported instruction file types inside the archive,
        extracts them to a temp directory, parses each one, and returns
        the concatenated text.
        """
        parseable = {".docx", ".pdf", ".txt", ".rtf", ".doc"}
        texts: list[str] = []

        with tempfile.TemporaryDirectory(prefix="excel_engine_zip_") as tmp:
            tmp_dir = Path(tmp)
            with zipfile.ZipFile(path, "r") as zf:
                for member in zf.namelist():
                    # Skip directories and hidden/system files
                    if member.endswith("/") or member.startswith("__MACOSX"):
                        continue
                    ext = Path(member).suffix.lower()
                    if ext not in parseable:
                        continue
                    zf.extract(member, tmp_dir)
                    extracted = tmp_dir / member
                    try:
                        texts.append(self.parse(extracted))
                    except Exception as e:
                        logger.warning("Failed to parse %s from ZIP: %s", member, e)

        if not texts:
            raise ValueError(
                f"No parseable instruction files found in ZIP: {path.name}. "
                f"Expected .docx, .pdf, .txt, or .rtf inside the archive."
            )
        return "\n\n".join(texts)

    # ── Helpers ──

    @staticmethod
    def _textutil_convert(path: Path) -> str:
        """
        Convert a document to plain text using macOS textutil.
        Works for .docx, .rtf, .rtfd, .doc, .html, .webarchive.
        """
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode != 0:
            raise RuntimeError(f"textutil failed: {result.stderr.strip()}")
        return result.stdout

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean up extracted text — normalize whitespace, remove artifacts."""
        # Collapse multiple blank lines
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Remove form feed and other control chars (keep \n, \t)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        return text.strip()

    # ── Structured SAM Instruction Parsing ──

    # Action verbs that commonly start SAM instruction lines
    _ACTION_VERBS = re.compile(
        r"^(Enter|Format|Create|Insert|Apply|Sort|Filter|Go\s+to|Navigate|"
        r"Select|In\s+cell|Use|Delete|Remove|Copy|Paste|Change|Set|Add|"
        r"Merge|Resize|Adjust|Move|Rename|Type|Save)\b",
        re.I,
    )

    # Patterns that mark step boundaries
    _NUMBERED_STEP = re.compile(r"^\s*(\d+)[.)]\s+(.+)", re.S)
    _PREFIXED_STEP = re.compile(r"^\s*Step\s+(\d+)[:.]\s*(.+)", re.I | re.S)
    _LETTERED_SUB = re.compile(r"^\s*([a-z])[.)]\s+(.+)", re.S)
    _BULLET = re.compile(r"^\s*[•\-▪●]\s+(.+)", re.S)

    # Sheet-switch detection
    _SHEET_SWITCH = re.compile(
        r"(?:Go\s+to|Switch\s+to|Navigate\s+to|On|In)\s+(?:the\s+)?"
        r"[\"']?([A-Za-z][\w\s\-]*?)[\"']?"
        r"\s+(?:work)?sheet\b",
        re.I,
    )

    # Table-name detection
    _TABLE_NAME = re.compile(
        r"(?:In|of|to|from|the)\s+(?:the\s+)?([A-Z][\w\s]*?)\s+table\b",
        re.I,
    )

    # Cross-reference patterns
    _CROSS_REF_REPEAT = re.compile(
        r"(?:repeat|do\s+the\s+same)\s+(?:for|with)\s+(?:column\s+)?([A-Z]{1,3}(?:\d+)?(?:\s*(?:through|to)\s*[A-Z]{1,3}\d*)?)",
        re.I,
    )
    _CROSS_REF_SAME = re.compile(
        r"(?:using|with|apply)\s+(?:the\s+)?same\s+(format|formula|style|function)",
        re.I,
    )
    _CROSS_REF_DO_SAME = re.compile(
        r"do\s+the\s+same\s+for\s+(?:cells?\s+)?([A-Z]{1,3}\d*\s*(?:through|to)\s*[A-Z]{1,3}\d*|[A-Z]{1,3}\d+)",
        re.I,
    )

    def split_into_steps(self, text: str) -> list[InstructionStep]:
        """
        Split raw instruction text into structured InstructionStep objects.

        Recognises numbered steps, lettered sub-steps, bullet points,
        and SAM-style action-verb lines.
        """
        lines = text.split("\n")
        steps: list[InstructionStep] = []
        step_num = 0
        current_parent: Optional[int] = None

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue

            # ── Numbered step: "1." / "2)" ──
            m = self._NUMBERED_STEP.match(line)
            if m:
                step_num += 1
                current_parent = step_num
                steps.append(InstructionStep(
                    step_number=step_num,
                    text=m.group(2).strip(),
                ))
                continue

            # ── Prefixed step: "Step 1:" ──
            m = self._PREFIXED_STEP.match(line)
            if m:
                step_num += 1
                current_parent = step_num
                steps.append(InstructionStep(
                    step_number=step_num,
                    text=m.group(2).strip(),
                ))
                continue

            # ── Lettered sub-step: "a.", "b)" ──
            m = self._LETTERED_SUB.match(line)
            if m:
                step_num += 1
                steps.append(InstructionStep(
                    step_number=step_num,
                    text=m.group(2).strip(),
                    parent_step=current_parent,
                ))
                continue

            # ── Bullet point ──
            m = self._BULLET.match(line)
            if m:
                step_num += 1
                steps.append(InstructionStep(
                    step_number=step_num,
                    text=m.group(1).strip(),
                    parent_step=current_parent,
                ))
                continue

            # ── SAM-style action verb at start of line ──
            if self._ACTION_VERBS.match(line) and len(line) > 10:
                step_num += 1
                current_parent = step_num
                steps.append(InstructionStep(
                    step_number=step_num,
                    text=line,
                ))
                continue

        return steps

    def carry_context(self, steps: list[InstructionStep]) -> list[InstructionStep]:
        """
        Propagate sheet and table context across sequential steps.

        If step N says "Go to the Revenue worksheet", all subsequent
        steps inherit "Revenue" as sheet_context until another sheet switch.
        Table name context is similarly propagated.
        """
        current_sheet: Optional[str] = None
        current_table: Optional[str] = None

        for step in steps:
            # Check for a sheet switch in this step
            sheet_match = self._SHEET_SWITCH.search(step.text)
            if sheet_match:
                current_sheet = sheet_match.group(1).strip()

            # Check for a table name reference in this step
            table_match = self._TABLE_NAME.search(step.text)
            if table_match:
                candidate = table_match.group(1).strip()
                # Filter out generic words that are not real table names
                if candidate.lower() not in ("the", "a", "an", "this", "that"):
                    current_table = candidate

            # Inherit the current sheet context
            if current_sheet and not step.sheet_context:
                step.sheet_context = current_sheet

            # Resolve "the table" references to the tracked table name
            if current_table and re.search(r"\bthe\s+table\b", step.text, re.I):
                step.text = re.sub(
                    r"\bthe\s+table\b",
                    f"the {current_table} table",
                    step.text,
                    count=1,
                    flags=re.I,
                )

        return steps

    def resolve_cross_references(
        self, steps: list[InstructionStep]
    ) -> list[InstructionStep]:
        """
        Detect and expand cross-reference phrases like "repeat for column D"
        or "using the same format".

        Clones the previous step's text with the new target when a cross-
        reference is detected, and tags the resulting step with a
        ``params.cross_ref`` marker in the text.
        """
        result: list[InstructionStep] = []
        prev_step: Optional[InstructionStep] = None

        for step in steps:
            # ── "repeat for column X" / "do the same for X through Y" ──
            repeat = self._CROSS_REF_REPEAT.search(step.text)
            do_same = self._CROSS_REF_DO_SAME.search(step.text)
            same_kind = self._CROSS_REF_SAME.search(step.text)

            if (repeat or do_same) and prev_step is not None:
                target = (repeat or do_same).group(1).strip()
                new_text = f"{prev_step.text} [cross-ref: apply to {target}]"
                result.append(InstructionStep(
                    step_number=step.step_number,
                    text=new_text,
                    sheet_context=step.sheet_context or prev_step.sheet_context,
                    parent_step=step.parent_step,
                ))
            elif same_kind and prev_step is not None:
                kind = same_kind.group(1)
                new_text = f"{step.text} [cross-ref: same {kind} as step {prev_step.step_number}]"
                result.append(InstructionStep(
                    step_number=step.step_number,
                    text=new_text,
                    sheet_context=step.sheet_context or prev_step.sheet_context,
                    parent_step=step.parent_step,
                ))
            else:
                result.append(step)

            prev_step = step  # always track the *original* step

        return result
